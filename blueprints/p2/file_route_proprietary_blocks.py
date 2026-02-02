from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from flask_login import login_required, current_user
from blueprints.p2.models import Folder, File, db
from datetime import datetime
from sqlalchemy.orm.attributes import flag_modified
import base64
import json
import io

from blueprints.p2.utils import save_data_uri_images_in_json
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

from . import combined_bp  # Import the blueprint instance


def format_file_size(size_bytes):
    """Format file size in human-readable format."""
    if size_bytes < 1024:
        return f"{size_bytes}B"
    elif size_bytes < 1024 * 1024:
        return f"{(size_bytes / 1024):.1f}KB"
    else:
        return f"{(size_bytes / (1024 * 1024)):.1f}MB"


def process_miobook_images(content_data, user_id):
    """
    Persist embedded data URI images to disk (deduped + compressed) and replace with URLs.
    Returns (updated_content_data, total_bytes_added).
    """
    total_bytes = 0

    if not content_data or not isinstance(content_data, dict):
        return content_data, total_bytes

    blocks = content_data.get('blocks', []) or []
    for block in blocks:
        if not isinstance(block, dict):
            continue

        try:
            updated_content, added = save_data_uri_images_in_json(block.get('content'), user_id)
            block['content'] = updated_content
            total_bytes += added
        except Exception as e:
            print(f"[WARN] Failed to process images for MioBook block {block.get('id')}: {e}")

    return content_data, total_bytes


def generate_default_miobook_title():
    """Return the default MioBook title using the shared timestamp convention."""
    return f"MioBook {datetime.now().strftime('%Y-%m-%d %H:%M')}"


def sanitize_filename_for_download(filename: str) -> str:
    """Sanitize filenames for downloads to avoid invalid characters and trim length."""
    invalid_chars = '<>:"/\\|?*'
    safe_name = filename or 'MioBook'
    for ch in invalid_chars:
        safe_name = safe_name.replace(ch, '_')
    return safe_name[:200]


def html_to_plain_text(html_content: str) -> str:
    """Convert rich HTML content into readable plain text for Word export."""
    if not html_content:
        return ''
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text('\n', strip=True)


def extract_board_image_bytes(block: dict):
    """Best-effort extraction of an image representation from a board block."""
    content = block.get('content') if isinstance(block, dict) else None
    data_url = None

    if isinstance(content, str) and content.startswith('data:image'):
        data_url = content
    elif isinstance(content, dict):
        # Prefer explicit imageData if provided
        image_data = content.get('imageData') or content.get('image')
        if isinstance(image_data, str) and image_data.startswith('data:image'):
            data_url = image_data

    if not data_url:
        return None

    try:
        header, encoded = data_url.split(',', 1)
        return base64.b64decode(encoded)
    except Exception as exc:  # noqa: PIE786
        print(f"[WARN] Failed to decode board image data: {exc}")
        return None


def add_markdown_block(export_doc: Document, block: dict):
    """Render a markdown block as plain paragraphs in Word export."""
    content = block.get('content', '') if isinstance(block, dict) else ''
    if not content:
        export_doc.add_paragraph('(No content)')
        return
    for line in str(content).splitlines():
        export_doc.add_paragraph(line)


def add_code_block(export_doc: Document, block: dict):
    """Render a code block with its language label."""
    language = ''
    metadata = block.get('metadata') if isinstance(block, dict) else None
    if isinstance(metadata, dict):
        language = metadata.get('language') or ''
    code_text = block.get('content', '') if isinstance(block, dict) else ''

    lang_label = f"Language: {language}" if language else 'Language: (unspecified)'
    export_doc.add_paragraph(lang_label)
    if code_text:
        export_doc.add_paragraph(str(code_text))
    else:
        export_doc.add_paragraph('(No code)')


def add_todo_block(export_doc: Document, block: dict):
    """Render todo items as a checklist list in Word export."""
    todos = block.get('content') if isinstance(block, dict) else None
    if not todos:
        export_doc.add_paragraph('(No todos)')
        return

    if isinstance(todos, list):
        for item in todos:
            if not isinstance(item, dict):
                export_doc.add_paragraph(f"- {item}")
                continue
            title = item.get('text') or item.get('title') or '(Untitled task)'
            done = item.get('checked') or item.get('done') or item.get('completed')
            prefix = '[x]' if done else '[ ]'
            export_doc.add_paragraph(f"{prefix} {title}")
    else:
        export_doc.add_paragraph(str(todos))


def add_editorjs_block(export_doc: Document, block: dict):
    """Render Editor.js content as bullet summaries."""
    content = block.get('content') if isinstance(block, dict) else None
    if not content:
        export_doc.add_paragraph('(No content)')
        return

    try:
        parsed = content if isinstance(content, dict) else json.loads(content)
    except Exception:
        export_doc.add_paragraph(str(content))
        return

    blocks = parsed.get('blocks') if isinstance(parsed, dict) else None
    if not blocks or not isinstance(blocks, list):
        export_doc.add_paragraph(str(parsed))
        return

    for editor_block in blocks:
        if not isinstance(editor_block, dict):
            export_doc.add_paragraph(str(editor_block))
            continue
        b_type = editor_block.get('type') or 'block'
        data = editor_block.get('data')
        summary = ''
        if isinstance(data, dict):
            if 'text' in data:
                summary = html_to_plain_text(data.get('text'))
            elif 'content' in data:
                summary = html_to_plain_text(data.get('content'))
            elif 'items' in data and isinstance(data['items'], list):
                summary = '; '.join([str(item) for item in data['items']])
            else:
                summary = json.dumps(data)
        else:
            summary = str(data) if data is not None else ''
        export_doc.add_paragraph(f"{b_type}: {summary or '(empty)'}")


def add_annotation_block(export_doc: Document, annotation: dict, index: int):
    """Render a single annotation (stored as its own block or inline on parent)."""
    if not isinstance(annotation, dict):
        export_doc.add_paragraph(f"Annotation {index}: {annotation}")
        return

    title = annotation.get('title') or ''
    if title:
        export_doc.add_paragraph(f"Annotation {index}: {title}")
    content = annotation.get('content')
    if not content:
        export_doc.add_paragraph('(No annotation content)')
        return

    # If annotation content is Editor.js structure, pull text from blocks
    def render_editorjs_content(raw):
        try:
            parsed = raw if isinstance(raw, dict) else json.loads(raw)
        except Exception:
            return None
        blocks = parsed.get('blocks') if isinstance(parsed, dict) else None
        if not blocks or not isinstance(blocks, list):
            return None
        texts = []
        for blk in blocks:
            if not isinstance(blk, dict):
                continue
            data = blk.get('data')
            if isinstance(data, dict) and 'text' in data:
                texts.append(html_to_plain_text(data.get('text')))
        if texts:
            for line in texts:
                export_doc.add_paragraph(line)
            return True
        return None

    rendered = render_editorjs_content(content)
    if rendered:
        return

    if isinstance(content, str):
        export_doc.add_paragraph(html_to_plain_text(content))
    else:
        export_doc.add_paragraph(str(content))


def extract_miobook_submission(request):
    """Extract user-provided fields for MioBook creation/update from JSON or form payloads."""
    payload = request.get_json(silent=True) if request.is_json else None

    if payload:
        raw_title = (payload.get('title') or '').strip()
        raw_folder_id = payload.get('folder_id')
        content_json_raw = payload.get('content_json', '{}')
    else:
        raw_title = (request.form.get('title', '') or '').strip()
        raw_folder_id = request.form.get('folder_id', type=int)
        content_json_raw = request.form.get('content_json', '{}')

    folder_id = None
    if raw_folder_id is not None:
        try:
            folder_id = int(raw_folder_id)
        except (TypeError, ValueError):
            folder_id = None

    return raw_title, folder_id, content_json_raw

@combined_bp.route('/new', methods=['GET', 'POST'])
@login_required
def new_combined():
    """Create a new MioBook document"""
    # Allow explicit folder targeting (e.g., Graph Workspace) for GET and POST.
    folder_override_id = (
        request.args.get('folder_id', type=int)
        or request.form.get('folder_id', type=int)
    )

    current_folder_id = session.get('current_folder_id')

    if folder_override_id:
        override_folder = Folder.query.filter_by(id=folder_override_id, user_id=current_user.id).first()
        if override_folder:
            current_folder_id = folder_override_id
            session['current_folder_id'] = current_folder_id
    
    if not current_folder_id:
        root_folder = Folder.query.filter_by(user_id=current_user.id, parent_id=None).first()
        current_folder_id = root_folder.id if root_folder else None

    if request.method == 'POST':
        try:
            raw_title, requested_folder_id, content_json_raw = extract_miobook_submission(request)
            title = raw_title or generate_default_miobook_title()
            folder_id = requested_folder_id or current_folder_id
            
            # Validate folder ownership
            if folder_id:
                valid_folder = Folder.query.filter_by(id=folder_id, user_id=current_user.id).first()
                if not valid_folder:
                    folder_id = current_folder_id

            # Get the combined content (new format with version)
            # Parse the combined content to validate it
            try:
                if isinstance(content_json_raw, dict):
                    content_data = content_json_raw
                else:
                    content_data = json.loads(content_json_raw)
                if not isinstance(content_data, dict):
                    content_data = {'version': '1.0', 'blocks': []}
                if 'version' not in content_data:
                    content_data['version'] = '1.0'
                if 'blocks' not in content_data or not isinstance(content_data['blocks'], list):
                    content_data['blocks'] = []
            except (TypeError, json.JSONDecodeError):
                content_data = {'version': '1.0', 'blocks': []}
            
            # Persist embedded images (e.g., whiteboard) to disk for dedupe/storage accounting
            content_data, bytes_added = process_miobook_images(content_data, current_user.id)

            # Store as a File with type='proprietary_blocks' and JSON content
            book_file = File(
                title=title,
                type='proprietary_blocks',
                content_json=content_data,
                folder_id=folder_id,
                owner_id=current_user.id,
                metadata_json={'description': 'MioBook combined document'}
            )
            
            # Calculate size and check limits
            def calculate_content_size(content):
                return len(content.encode('utf-8')) if content else 0
            
            def check_guest_limit(user, additional_size):
                if getattr(user, 'user_type', None) == 'guest':
                    max_size = 50 * 1024 * 1024
                    if (user.total_data_size or 0) + additional_size > max_size:
                        flash("Data limit exceeded (50MB max for guests). Please delete some data or upgrade your account.", "danger")
                        return False
                return True
            
            def update_user_data_size(user, delta):
                user.total_data_size = (user.total_data_size or 0) + delta
                db.session.commit()
            
            content_size = book_file.get_content_size()
            total_size = content_size + (bytes_added or 0)

            if not check_guest_limit(current_user, total_size):
                return jsonify({"success": False, "error": "Data limit exceeded"}), 400
            
            # Save to database
            db.session.add(book_file)
            db.session.commit()
            update_user_data_size(current_user, total_size)
            
            # Add notification for creation
            from blueprints.p2.utils import add_notification
            size_str = format_file_size(content_size)
            notif_msg = f"Created MioBook '{title}' ({size_str})"
            add_notification(current_user.id, notif_msg, 'save')
            
            # Check if this is an AJAX request or JSON request
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
                return jsonify({"success": True, "document_id": book_file.id})
            
            flash("MioBook created successfully!", "success")
            return redirect(url_for('combined.edit_combined', document_id=book_file.id))
            
        except Exception as e:
            print(f"[ERROR] Failed to create MioBook: {e}")
            db.session.rollback()
            flash("Failed to create MioBook. Please try again.", "danger")
            return redirect(url_for('folders.view_folder', folder_id=current_folder_id))
    
    # GET request - show creation page
    folders = Folder.query.filter_by(user_id=current_user.id).all()
    current_folder = Folder.query.filter_by(id=current_folder_id, user_id=current_user.id).first() if current_folder_id else None
    
    return render_template('p2/file_edit_proprietary_blocks.html', 
                         folders=folders, 
                         current_folder=current_folder,
                         folder_id=current_folder_id,
                         default_title=generate_default_miobook_title())

@combined_bp.route('/edit/<int:document_id>', methods=['GET', 'POST'])
@login_required
def edit_combined(document_id):
    """Edit an existing MioBook document"""
    # Fetch the File record with type='proprietary_blocks'
    document = File.query.filter_by(id=document_id, owner_id=current_user.id, type='proprietary_blocks').first()
    if not document:
        flash("Document not found or access denied.", "danger")
        return redirect(url_for('folders.view_folder', folder_id=session.get('current_folder_id')))
    
    if request.method == 'POST':
        try:
            # Update title
            new_title = request.form.get('title', '').strip() if not request.is_json else request.get_json().get('title', '').strip()
            if new_title:
                document.title = new_title
            
            # Update combined content (new format with version)
            content_json_str = request.get_json().get('content_json', '{}') if request.is_json else request.form.get('content_json', '{}')
            
            # Parse the combined content to validate it
            try:
                content_data = json.loads(content_json_str)
                if not isinstance(content_data, dict):
                    content_data = {'version': '1.0', 'blocks': []}
                if 'version' not in content_data:
                    content_data['version'] = '1.0'
                if 'blocks' not in content_data or not isinstance(content_data['blocks'], list):
                    content_data['blocks'] = []
                
                # Debug: Log what we're saving
                print(f"[DEBUG] Saving MioBook {document_id}: '{new_title or document.title}'")
                print(f"[DEBUG] Number of blocks to save: {len(content_data['blocks'])}")
                for i, block in enumerate(content_data['blocks']):
                    print(f"[DEBUG] Block {i}: type={block.get('type')}, id={block.get('id')}, title={block.get('title')}")
                    if block.get('type') == 'whiteboard':
                        content = block.get('content')
                        print(f"[DEBUG] Whiteboard content type: {type(content)}")
                        if isinstance(content, dict):
                            print(f"[DEBUG] Whiteboard content keys: {content.keys()}")
                            if 'imageData' in content:
                                img_data = content['imageData']
                                print(f"[DEBUG] imageData present: {bool(img_data)}, length: {len(img_data) if img_data else 0}")
                        else:
                            print(f"[DEBUG] Whiteboard content: {str(content)[:100] if content else 'None/Empty'}")
                
            except json.JSONDecodeError:
                content_data = {'version': '1.0', 'blocks': []}
            
            # Persist embedded images (e.g., whiteboard) to disk for dedupe/storage accounting
            content_data, bytes_added = process_miobook_images(content_data, current_user.id)

            # Calculate size difference
            old_size = document.get_content_size()
            document.content_json = content_data
            flag_modified(document, 'content_json')  # Required for SQLAlchemy to detect JSON changes
            new_size = document.get_content_size()
            size_delta = (new_size - old_size) + (bytes_added or 0)
            
            # Check guest limits for size increase
            def check_guest_limit(user, additional_size):
                if getattr(user, 'user_type', None) == 'guest' and additional_size > 0:
                    max_size = 50 * 1024 * 1024
                    if (user.total_data_size or 0) + additional_size > max_size:
                        flash("Data limit exceeded (50MB max for guests). Please delete some data or upgrade your account.", "danger")
                        return False
                return True
            
            def update_user_data_size(user, delta):
                user.total_data_size = (user.total_data_size or 0) + delta
                db.session.commit()
            
            if not check_guest_limit(current_user, size_delta):
                return jsonify({"success": False, "error": "Data limit exceeded"}), 400
            
            # Save changes
            document.last_modified = datetime.utcnow()
            db.session.commit()
            update_user_data_size(current_user, size_delta)
            
            # Add notification for save
            from blueprints.p2.utils import add_notification
            size_str = format_file_size(new_size)
            notif_msg = f"Saved MioBook '{document.title}' ({size_str})"
            add_notification(current_user.id, notif_msg, 'save')
            
            # Check if this is an AJAX request or JSON request
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
                return jsonify({"success": True})
            
            # Use telemetry notification instead of flash
            return redirect(url_for('combined.edit_combined', document_id=document_id, saved='miobook', size=size_str))
            
        except Exception as e:
            print(f"[ERROR] Failed to update MioBook: {e}")
            db.session.rollback()
            flash("Failed to update MioBook. Please try again.", "danger")
    
    # Document already has content_json in the new format
    # Template will handle parsing it
    folders = Folder.query.filter_by(user_id=current_user.id).all()
    current_folder = document.folder
    
    # Debug: Log the document content structure
    print(f"[DEBUG] Loading MioBook {document_id}: '{document.title}'")
    print(f"[DEBUG] content_json type: {type(document.content_json)}")
    if document.content_json:
        print(f"[DEBUG] content_json keys: {document.content_json.keys() if isinstance(document.content_json, dict) else 'not a dict'}")
        if isinstance(document.content_json, dict) and 'blocks' in document.content_json:
            print(f"[DEBUG] Number of blocks: {len(document.content_json['blocks'])}")
            for i, block in enumerate(document.content_json['blocks']):
                print(f"[DEBUG] Block {i}: type={block.get('type')}, id={block.get('id')}, title={block.get('title')}")
                if block.get('type') == 'whiteboard':
                    content = block.get('content')
                    print(f"[DEBUG] Whiteboard content type: {type(content)}")
                    if isinstance(content, dict):
                        print(f"[DEBUG] Whiteboard content keys: {content.keys()}")
                        if 'imageData' in content:
                            img_data = content['imageData']
                            print(f"[DEBUG] imageData present: {bool(img_data)}, length: {len(img_data) if img_data else 0}")
                    else:
                        print(f"[DEBUG] Whiteboard content value: {str(content)[:100] if content else 'None/Empty'}")
    
    return render_template('p2/file_edit_proprietary_blocks.html', 
                         file=document,  # Pass as 'file' to match template expectations
                         document=document,  # Keep for backward compatibility
                         folders=folders, 
                         current_folder=current_folder,
                         folder_id=current_folder.id if current_folder else None,
                         default_title=document.title or generate_default_miobook_title())


# Legacy save_note_block and save_board_block routes removed
# MioBook now stores all blocks inline in content_json (no separate File records)
# Blocks are self-contained within the MioBook File record


@combined_bp.route('/download_json/<int:document_id>', methods=['GET'])
@login_required
def download_json(document_id):
    """Download combined document as JSON"""
    document = File.query.filter_by(id=document_id, owner_id=current_user.id, type='proprietary_blocks').first()
    if not document:
        flash("Document not found or access denied.", "danger")
        return redirect(url_for('folders.view_folder', folder_id=session.get('current_folder_id')))
    
    try:
        # Get content from content_json (new format with version)
        content_data = document.content_json if document.content_json else {'version': '1.0', 'blocks': []}
        if not isinstance(content_data, dict):
            content_data = {'version': '1.0', 'blocks': []}
        
        # Create export payload
        payload = {
            "title": document.title or "Untitled",
            "document_id": document.id,
            "exported_at": datetime.now().isoformat(),
            "type": "miobook",
            "version": content_data.get('version', '1.0'),
            "blocks": content_data.get('blocks', [])
        }
        
        # Create JSON file in memory
        json_data = json.dumps(payload, indent=2)
        json_buffer = io.BytesIO(json_data.encode('utf-8'))
        json_buffer.seek(0)
        
        # Safe filename
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in (document.title or 'Untitled'))
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        return send_file(
            json_buffer,
            mimetype='application/json',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"[ERROR] Failed to download JSON: {e}")
        flash("Failed to download document as JSON.", "danger")
        return redirect(url_for('combined.edit_combined', document_id=document_id))


@combined_bp.route('/download_docx/<int:document_id>', methods=['GET'])
@login_required
def download_docx(document_id):
    """Export a MioBlocks (proprietary_blocks) document to a Word .docx file."""
    document = File.query.filter_by(id=document_id, owner_id=current_user.id, type='proprietary_blocks').first()
    if not document:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'error': 'Document not found or access denied.'}), 404
        flash("Document not found or access denied.", "danger")
        return redirect(url_for('folders.view_folder', folder_id=session.get('current_folder_id')))

    try:
        content_data = document.content_json if document.content_json else {'version': '1.0', 'blocks': []}
        if not isinstance(content_data, dict):
            content_data = {'version': '1.0', 'blocks': []}

        blocks = content_data.get('blocks', []) or []

        annotations_map = {}
        main_blocks = []
        for blk in blocks:
            blk_type = blk.get('type') if isinstance(blk, dict) else None
            if blk_type == 'annotation':
                parent_id = blk.get('parentId') if isinstance(blk, dict) else None
                if parent_id is not None:
                    annotations_map.setdefault(parent_id, []).append(blk)
            else:
                main_blocks.append(blk)

        export_doc = Document()
        export_doc.add_heading(document.title or 'MioBook', 0)
        export_doc.add_paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        if not main_blocks:
            export_doc.add_paragraph("This MioBlocks document has no blocks yet.")

        for idx, block in enumerate(main_blocks, start=1):
            if idx > 1:
                export_doc.add_page_break()
            block_type = block.get('type') if isinstance(block, dict) else None
            export_doc.add_heading(f"Block {idx} - {block_type or 'Unknown'}", level=1)

            if block_type == 'note':
                plain_text = html_to_plain_text(block.get('content', '')) if isinstance(block, dict) else ''
                export_doc.add_paragraph(plain_text or '(No content)')
            elif block_type == 'markdown':
                add_markdown_block(export_doc, block)
            elif block_type == 'code':
                add_code_block(export_doc, block)
            elif block_type == 'todo':
                add_todo_block(export_doc, block)
            elif block_type == 'blocks':
                add_editorjs_block(export_doc, block)
            elif block_type in ('board', 'board-iframe', 'whiteboard'):
                image_bytes = extract_board_image_bytes(block) if isinstance(block, dict) else None
                if image_bytes:
                    try:
                        export_doc.add_picture(io.BytesIO(image_bytes), width=Inches(6))
                    except Exception as img_exc:  # noqa: PIE786
                        print(f"[WARN] Failed to embed board image: {img_exc}")
                        export_doc.add_paragraph('Whiteboard image could not be embedded, but is available in the app.')
                else:
                    export_doc.add_paragraph('Whiteboard content available in the app; image export not available in this file.')
            else:
                fallback = block.get('content') if isinstance(block, dict) else None
                export_doc.add_paragraph(str(fallback) if fallback else 'Unsupported block type for Word export.')

            annotations = []
            block_id = block.get('id') if isinstance(block, dict) else None

            if isinstance(block, dict):
                inline_annotations = block.get('annotations')
                if isinstance(inline_annotations, list):
                    annotations.extend(inline_annotations)

            if block_id is not None and block_id in annotations_map:
                annotations.extend(annotations_map.get(block_id) or [])

            if annotations:
                export_doc.add_heading('Annotations', level=2)
                for a_idx, annotation in enumerate(annotations, start=1):
                    add_annotation_block(export_doc, annotation, a_idx)

        output = io.BytesIO()
        export_doc.save(output)
        output.seek(0)

        filename = f"{sanitize_filename_for_download(document.title)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as exc:  # noqa: PIE786
        print(f"[ERROR] Failed to export MioBook to DOCX: {exc}")
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'error': 'Failed to export document.'}), 500
        flash("Failed to export document as Word.", "danger")
        return redirect(url_for('combined.edit_combined', document_id=document_id))


@combined_bp.route('/print_view/<int:document_id>', methods=['GET'])
@login_required
def print_view(document_id):
    """Open print-friendly view for PDF generation via browser"""
    document = File.query.filter_by(id=document_id, owner_id=current_user.id, type='proprietary_blocks').first()
    if not document:
        flash("Document not found or access denied.", "danger")
        return redirect(url_for('folders.view_folder', folder_id=session.get('current_folder_id')))
    
    try:
        # Get content from content_json (new format with version)
        content_data = document.content_json if document.content_json else {'version': '1.0', 'blocks': []}
        if not isinstance(content_data, dict):
            content_data = {'version': '1.0', 'blocks': []}
        content_blocks = content_data.get('blocks', [])
        
        # Render print view template
        return render_template('p2/miobook_print_view.html', 
                             document=document,
                             content_blocks=content_blocks)
            
    except Exception as e:
        print(f"[ERROR] Failed to open print view: {e}")
        flash("Failed to open print view.", "danger")
        return redirect(url_for('combined.edit_combined', document_id=document_id))

