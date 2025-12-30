


######################################
# notes , whiteboard side utilities
######################################
from datetime import datetime
from extensions import db, login_manager
from flask import flash
import os   



# build and pass the breadcrumb
def build_folder_breadcrumb(folder):
    """Return list of folders from root -> ... -> current folder."""
    chain = []
    f = folder
    while f is not None:
        chain.append(f)
        f = f.parent  # assumes Folder.parent relationship
    chain.reverse()
    print(chain)
    return chain

# --- Data size cap helpers ---
def calculate_content_size(content):
    return len(content.encode('utf-8')) if content else 0

def calculate_image_size(file_path):
    return os.path.getsize(file_path) if os.path.exists(file_path) else 0

# Check if guest user exceeds 50MB cap
def check_guest_limit(user, additional_size):
    if getattr(user, 'user_type', None) == 'guest':
        max_size = 50 * 1024 * 1024
        if (user.total_data_size or 0) + additional_size > max_size:
            flash("Data limit exceeded (50MB max for guests). Please delete some data or upgrade your account.", "danger")
            return False
    return True

# Update user's total data size
def update_user_data_size(user, delta):
    user.total_data_size = (user.total_data_size or 0) + delta
    db.session.commit()


######################################
# Document Parsing for Chat Attachments (Phase 2)
######################################

import hashlib
import chardet
from pathlib import Path
from werkzeug.utils import secure_filename

# PDF parsing
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# XLSX parsing
try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# OCR
try:
    import pytesseract
    from PIL import Image
    import config
    if hasattr(config, 'TESSERACT_CMD'):
        pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def calculate_file_hash(file_path):
    """Calculate SHA256 hash of file for deduplication"""
    sha256 = hashlib.sha256()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            sha256.update(chunk)
    return sha256.hexdigest()


def detect_file_encoding(file_path):
    """Detect character encoding of text file"""
    with open(file_path, 'rb') as f:
        raw_data = f.read(100000)  # Sample first 100KB
        result = chardet.detect(raw_data)
        return result['encoding'] or 'utf-8'


def parse_pdf_to_text(file_path):
    """Extract text from PDF file"""
    if not PYPDF_AVAILABLE:
        raise ImportError("pypdf not installed. Run: pip install pypdf")
    
    reader = PdfReader(file_path)
    text_parts = []
    
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            text_parts.append(f"--- Page {page_num} ---\n{text}")
    
    return "\n\n".join(text_parts)


def parse_docx_to_text(file_path):
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = Document(file_path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    
    return "\n\n".join(text_parts)


def parse_excel_to_text(file_path):
    """Extract data from XLSX file as structured text"""
    if not OPENPYXL_AVAILABLE:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    
    wb = load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"=== Sheet: {sheet_name} ===")
        
        rows_data = []
        for row in sheet.iter_rows(values_only=True):
            # Filter out completely empty rows
            if any(cell is not None for cell in row):
                row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                rows_data.append(row_text)
        
        text_parts.append("\n".join(rows_data))
    
    return "\n\n".join(text_parts)


def extract_text_from_image_ocr(file_path):
    """Extract text from image using OCR"""
    if not OCR_AVAILABLE:
        raise ImportError("pytesseract not installed or Tesseract not found")
    
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text.strip()


def read_text_file_safely(file_path, max_size_mb=10):
    """Read text file with encoding detection and size limit"""
    file_size = os.path.getsize(file_path)
    max_size = max_size_mb * 1024 * 1024
    
    if file_size > max_size:
        raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f}MB (max {max_size_mb}MB)")
    
    encoding = detect_file_encoding(file_path)
    
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 which never fails
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def parse_document_for_chat(file_path, file_type):
    """
    Central dispatcher for document parsing
    Returns tuple: (text_content, parse_method)
    """
    file_path = str(file_path)
    
    if file_type == 'pdf':
        return parse_pdf_to_text(file_path), 'pdf_extraction'
    
    elif file_type in ['docx', 'doc']:
        return parse_docx_to_text(file_path), 'docx_extraction'
    
    elif file_type in ['xlsx', 'xls']:
        return parse_excel_to_text(file_path), 'excel_extraction'
    
    elif file_type == 'image':
        return extract_text_from_image_ocr(file_path), 'ocr'
    
    elif file_type in ['py', 'js', 'html', 'css', 'ts', 'yaml', 'yml', 'env', 'json', 'md', 'txt']:
        return read_text_file_safely(file_path), 'text_read'
    
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def get_file_type_from_extension(filename):
    """Map file extension to internal file type identifier"""
    ext = Path(filename).suffix.lower().lstrip('.')
    
    # Image types
    if ext in ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp']:
        return 'image'
    
    # Map extensions to types
    type_map = {
        'pdf': 'pdf',
        'docx': 'docx',
        'doc': 'doc',
        'xlsx': 'xlsx',
        'xls': 'xls',
        'py': 'py',
        'js': 'js',
        'ts': 'ts',
        'html': 'html',
        'css': 'css',
        'yaml': 'yaml',
        'yml': 'yaml',
        'env': 'env',
        'json': 'json',
        'md': 'md',
        'txt': 'txt',
    }
    
    return type_map.get(ext, 'unknown')
